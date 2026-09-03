from __future__ import annotations

import json
import math
import textwrap
from io import StringIO
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as RLImage,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from src.eda_utils import (
    PROJECT_ROOT,
    balanced_label,
    dataframe_profile,
    load_dataset,
    missing_summary,
    normality_summary,
    outlier_summary,
    statistical_summary,
    target_distribution,
    target_percentages,
    encode_for_correlation,
)

NOTEBOOK_PATH = PROJECT_ROOT / "notebooks" / "01_EDA.ipynb"
REPORTS_DIR = PROJECT_ROOT / "reports"
FIGURES_DIR = PROJECT_ROOT / "figures"
LAB_MD = REPORTS_DIR / "ML_Lab_Report.md"
LAB_DOCX = REPORTS_DIR / "ML_Lab_Report.docx"
LAB_PDF = REPORTS_DIR / "ML_Lab_Report.pdf"
OUTPUT_IMAGE_DIR = FIGURES_DIR


def ensure_directories() -> None:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    FIGURES_DIR.mkdir(parents=True, exist_ok=True)


def load_notebook_code_cells() -> list[str]:
    notebook = json.loads(NOTEBOOK_PATH.read_text(encoding="utf-8"))
    code_cells: list[str] = []
    for cell in notebook["cells"]:
        if cell["cell_type"] == "code":
            code_cells.append("".join(cell["source"]))
    return code_cells


def extract_notebook_cell_sources() -> dict[str, str]:
    code_cells = load_notebook_code_cells()
    return {
        "q1_program": code_cells[0] + "\n\n" + code_cells[1],
        "q2_program": code_cells[2],
        "q3_program": code_cells[3],
        "q4_program": code_cells[4],
        "q5_program": code_cells[5],
        "q6_program": code_cells[6],
        "q7_program": code_cells[7] + "\n\n" + code_cells[8],
        "q8_program": code_cells[9],
        "q9_program": code_cells[10],
    }


def wrap_text(text: str, width: int = 110) -> str:
    lines: list[str] = []
    for paragraph in text.splitlines():
        if not paragraph.strip():
            lines.append("")
        else:
            lines.extend(textwrap.wrap(paragraph, width=width) or [""])
    return "\n".join(lines)


def render_text_image(text: str, output_path: Path, title: str) -> Path:
    plt.figure(figsize=(12, max(2.5, 0.42 * text.count("\n") + 2)))
    ax = plt.gca()
    ax.axis("off")
    ax.set_title(title, fontsize=14, fontweight="bold", pad=12)
    ax.text(
        0.01,
        0.98,
        wrap_text(text, width=120),
        va="top",
        ha="left",
        family="monospace",
        fontsize=10,
        transform=ax.transAxes,
        bbox={"facecolor": "white", "edgecolor": "#999999", "boxstyle": "round,pad=0.6"},
    )
    plt.savefig(output_path, dpi=300, bbox_inches="tight")
    plt.close()
    return output_path


def create_text_output_images(df: pd.DataFrame) -> dict[str, Path]:
    output_paths: dict[str, Path] = {}
    output_paths["q1_head"] = render_text_image(df.head().to_string(), OUTPUT_IMAGE_DIR / "q1_head_output.png", "Question 1 - head()")
    output_paths["q1_tail"] = render_text_image(df.tail().to_string(), OUTPUT_IMAGE_DIR / "q1_tail_output.png", "Question 1 - tail()")
    output_paths["q1_shape_columns"] = render_text_image(
        f"Shape: {df.shape}\n\nColumns:\n{list(df.columns)}",
        OUTPUT_IMAGE_DIR / "q1_shape_columns_output.png",
        "Question 1 - Shape and Columns",
    )
    output_paths["q1_dtypes"] = render_text_image(df.dtypes.to_string(), OUTPUT_IMAGE_DIR / "q1_dtypes_output.png", "Question 1 - Data Types")

    info_buffer = StringIO()
    df.info(buf=info_buffer)
    output_paths["q1_info"] = render_text_image(info_buffer.getvalue(), OUTPUT_IMAGE_DIR / "q1_info_output.png", "Question 1 - info()")

    output_paths["q2_describe_numeric"] = render_text_image(
        df.describe().round(3).to_string(),
        OUTPUT_IMAGE_DIR / "q2_describe_numeric_output.png",
        "Question 2 - describe()",
    )
    output_paths["q2_describe_object"] = render_text_image(
        df.describe(include="object").to_string(),
        OUTPUT_IMAGE_DIR / "q2_describe_object_output.png",
        "Question 2 - describe(include='object')",
    )
    missing = missing_summary(df)
    output_paths["q2_missing_values"] = render_text_image(
        missing.to_string(),
        OUTPUT_IMAGE_DIR / "q2_missing_values_output.png",
        "Question 2 - Missing Values",
    )
    output_paths["q2_duplicates"] = render_text_image(
        f"Duplicate rows: {int(df.duplicated().sum())}",
        OUTPUT_IMAGE_DIR / "q2_duplicate_count_output.png",
        "Question 2 - Duplicate Count",
    )
    return output_paths


def create_missing_heatmap(df: pd.DataFrame) -> Path:
    output_path = FIGURES_DIR / "missing_values_heatmap.png"
    if output_path.exists():
        return output_path

    missing_percentages = (df.isna().sum() / len(df) * 100).to_frame(name="missing %")
    plt.figure(figsize=(12, 5))
    ax = sns.heatmap(missing_percentages.T, annot=True, fmt=".2f", cmap="magma", cbar=False, linewidths=0.3)
    ax.set_title("Missing Value Heatmap")
    ax.set_xlabel("Column")
    ax.set_ylabel("")
    plt.xticks(rotation=30, ha="right")
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches="tight")
    plt.close()
    return output_path


def get_figure_files() -> list[Path]:
    figure_files = sorted(FIGURES_DIR.glob("*.png"))
    return figure_files


def code_block(text: str) -> str:
    return "```python\n" + text.rstrip() + "\n```"


def markdown_table(headers: list[str], rows: list[list[str]]) -> str:
    header_line = "| " + " | ".join(headers) + " |"
    separator = "| " + " | ".join(["---"] * len(headers)) + " |"
    body = ["| " + " | ".join(row) + " |" for row in rows]
    return "\n".join([header_line, separator, *body])


def correlation_tables(df: pd.DataFrame) -> tuple[str, pd.DataFrame]:
    encoded = encode_for_correlation(df)
    corr = encoded.corr(numeric_only=True)
    stroke_corr = corr["stroke"].sort_values(ascending=False)
    top_positive = stroke_corr.iloc[1:6]
    top_negative = stroke_corr.tail(5).sort_values()
    rows = [[index, f"{value:.4f}"] for index, value in stroke_corr.items()]
    table_md = markdown_table(["Feature", "Correlation with Stroke"], rows)
    return table_md, corr


def format_question_block(title: str, aim: str, program: str, output: str, analysis: str, conclusion: str) -> str:
    return "\n".join(
        [
            f"## {title}",
            "",
            "### Aim",
            aim,
            "",
            "### Program",
            code_block(program),
            "",
            "### Output",
            output,
            "",
            "### Analysis",
            analysis,
            "",
            "### Conclusion",
            conclusion,
            "",
        ]
    )


def build_markdown_report(df: pd.DataFrame, code_cells: dict[str, str], output_images: dict[str, Path], corr: pd.DataFrame) -> str:
    profile = dataframe_profile(df)
    missing = missing_summary(df)
    stats = statistical_summary(df)
    normality = normality_summary(df)
    outliers = outlier_summary(df)
    target_counts = target_distribution(df)
    target_pct = target_percentages(df)
    corr_table_md, _ = correlation_tables(df)

    q1_output = "\n\n".join(
        [
            f"![head]({output_images['q1_head'].relative_to(PROJECT_ROOT).as_posix()})",
            f"![tail]({output_images['q1_tail'].relative_to(PROJECT_ROOT).as_posix()})",
            f"![shape and columns]({output_images['q1_shape_columns'].relative_to(PROJECT_ROOT).as_posix()})",
            f"![data types]({output_images['q1_dtypes'].relative_to(PROJECT_ROOT).as_posix()})",
            f"![info]({output_images['q1_info'].relative_to(PROJECT_ROOT).as_posix()})",
        ]
    )

    q2_output = "\n\n".join(
        [
            f"![describe numeric]({output_images['q2_describe_numeric'].relative_to(PROJECT_ROOT).as_posix()})",
            f"![describe object]({output_images['q2_describe_object'].relative_to(PROJECT_ROOT).as_posix()})",
            f"![missing values]({output_images['q2_missing_values'].relative_to(PROJECT_ROOT).as_posix()})",
            f"![duplicate count]({output_images['q2_duplicates'].relative_to(PROJECT_ROOT).as_posix()})",
        ]
    )

    q3_output = "\n\n".join(
        [
            f"![missing values heatmap]({(FIGURES_DIR / 'missing_values_heatmap.png').relative_to(PROJECT_ROOT).as_posix()})",
            f"![missing values table]({output_images['q2_missing_values'].relative_to(PROJECT_ROOT).as_posix()})",
            f"![duplicate count]({output_images['q2_duplicates'].relative_to(PROJECT_ROOT).as_posix()})",
        ]
    )

    q4_output = "\n\n".join(
        [
            f"![stroke countplot]({(FIGURES_DIR / 'stroke_distribution.png').relative_to(PROJECT_ROOT).as_posix()})",
            f"![stroke pie chart]({(FIGURES_DIR / 'stroke_pie_chart.png').relative_to(PROJECT_ROOT).as_posix()})",
            f"![stroke percentage distribution]({(FIGURES_DIR / 'stroke_percentage_distribution.png').relative_to(PROJECT_ROOT).as_posix()})",
            markdown_table(
                ["Stroke", "Count", "Percentage"],
                [[str(index), str(int(count)), f"{target_pct.loc[index]:.2f}%"] for index, count in target_counts.items()],
            ),
        ]
    )

    q5_output = "\n\n".join(
        [
            f"![age histogram]({(FIGURES_DIR / 'age_distribution.png').relative_to(PROJECT_ROOT).as_posix()})",
            f"![bmi histogram]({(FIGURES_DIR / 'bmi_distribution.png').relative_to(PROJECT_ROOT).as_posix()})",
            f"![glucose histogram]({(FIGURES_DIR / 'avg_glucose_level_distribution.png').relative_to(PROJECT_ROOT).as_posix()})",
            f"![age boxplot]({(FIGURES_DIR / 'age_boxplot.png').relative_to(PROJECT_ROOT).as_posix()})",
            f"![bmi boxplot]({(FIGURES_DIR / 'bmi_boxplot.png').relative_to(PROJECT_ROOT).as_posix()})",
            f"![glucose boxplot]({(FIGURES_DIR / 'avg_glucose_level_boxplot.png').relative_to(PROJECT_ROOT).as_posix()})",
        ]
    )

    q6_output = "\n\n".join(
        [
            f"![gender countplot]({(FIGURES_DIR / 'gender_countplot.png').relative_to(PROJECT_ROOT).as_posix()})",
            f"![work type countplot]({(FIGURES_DIR / 'work_type_countplot.png').relative_to(PROJECT_ROOT).as_posix()})",
            f"![ever married countplot]({(FIGURES_DIR / 'ever_married_countplot.png').relative_to(PROJECT_ROOT).as_posix()})",
            f"![residence type countplot]({(FIGURES_DIR / 'Residence_type_countplot.png').relative_to(PROJECT_ROOT).as_posix()})",
            f"![smoking status countplot]({(FIGURES_DIR / 'smoking_status_countplot.png').relative_to(PROJECT_ROOT).as_posix()})",
            f"![hypertension countplot]({(FIGURES_DIR / 'hypertension_countplot.png').relative_to(PROJECT_ROOT).as_posix()})",
            f"![heart disease countplot]({(FIGURES_DIR / 'heart_disease_countplot.png').relative_to(PROJECT_ROOT).as_posix()})",
        ]
    )

    q7_output = "\n\n".join(
        [
            f"![correlation matrix]({output_images.get('correlation_matrix_image', FIGURES_DIR / 'correlation_heatmap.png').relative_to(PROJECT_ROOT).as_posix()})",
            f"![correlation heatmap]({(FIGURES_DIR / 'correlation_heatmap.png').relative_to(PROJECT_ROOT).as_posix()})",
            corr_table_md,
        ]
    )

    q8_output = "\n\n".join(
        [
            f"![age outlier boxplot]({(FIGURES_DIR / 'age_outlier_boxplot.png').relative_to(PROJECT_ROOT).as_posix()})",
            f"![bmi outlier boxplot]({(FIGURES_DIR / 'bmi_outlier_boxplot.png').relative_to(PROJECT_ROOT).as_posix()})",
            f"![glucose outlier boxplot]({(FIGURES_DIR / 'avg_glucose_level_outlier_boxplot.png').relative_to(PROJECT_ROOT).as_posix()})",
        ]
    )

    final_analysis = [
        "- Dataset size: 5,110 rows and 12 columns.",
        "- Numerical features: age, avg_glucose_level, and bmi.",
        "- Categorical features: gender, work_type, ever_married, Residence_type, and smoking_status.",
        "- Missing values are concentrated in bmi, with 201 missing records.",
        "- No duplicate rows were found in the dataset.",
        "- The stroke target is highly imbalanced, with 95.13% class 0 and 4.87% class 1.",
        "- Age is the strongest feature associated with stroke.",
        "- Heart disease, hypertension, and average glucose level show positive association with stroke.",
        "- Former smokers and self-employed patients have higher stroke proportions.",
        "- Married patients show a higher stroke rate than unmarried patients.",
        "- BMI and average glucose level are right-skewed and contain visible outliers.",
        "- Residence type and gender show comparatively weak linear correlation with stroke.",
        "- The dataset quality is good overall, but preprocessing is needed before modeling.",
    ]

    report = [
        "# Machine Learning Lab Report: Stroke Prediction Dataset",
        "",
        "## Introduction",
        "This lab report documents the exploratory analysis of the Kaggle stroke prediction dataset. The objective is to understand the dataset structure, quality, distributions, relationships, and preparation needs before building a classification model.",
        "",
        "## Question 1",
        "### Aim",
        "Import, load and view the Stroke Prediction dataset.",
        "",
        "### Program",
        code_block(code_cells["q1_program"]),
        "",
        "### Output",
        q1_output,
        "",
        "### Analysis",
        f"The dataset loads successfully from the workspace. It has {profile['shape'][0]} rows and {profile['shape'][1]} columns, with the expected stroke-related fields and a small number of data type groups.",
        "",
        "### Conclusion",
        "The dataset is accessible and ready for analysis.",
        "",
        "## Question 2",
        "### Aim",
        "Display the summary statistics of the dataset.",
        "",
        "### Program",
        code_block(code_cells["q2_program"]),
        "",
        "### Output",
        q2_output,
        "",
        "### Analysis",
        "The summary statistics show that age is widely spread, average glucose level is strongly right-skewed, and BMI contains missing values. The categorical summary confirms the dominant categories in gender, work type, residence type, and smoking status.",
        "",
        "### Conclusion",
        "The descriptive statistics confirm that the dataset is realistic, imbalanced, and suitable for a detailed EDA phase.",
        "",
        "## Question 3",
        "### Aim",
        "Analyze missing values and duplicate records.",
        "",
        "### Program",
        code_block(code_cells["q3_program"]),
        "",
        "### Output",
        q3_output,
        "",
        "### Analysis",
        f"Only BMI contains missing values, accounting for {int(profile['missing_values'].sum())} missing entries overall. The heatmap makes the missingness pattern clear, and there are {profile['duplicate_rows']} duplicate rows, so the dataset remains structurally clean.",
        "",
        "### Conclusion",
        "The dataset requires missing-value handling for BMI, but duplicate removal is not necessary.",
        "",
        "## Question 4",
        "### Aim",
        "Analyze the target variable (Stroke).",
        "",
        "### Program",
        code_block(code_cells["q4_program"]),
        "",
        "### Output",
        q4_output,
        "",
        "### Analysis",
        f"Stroke cases are rare: class 0 has {target_counts.loc[0]} records ({target_pct.loc[0]:.2f}%) and class 1 has {target_counts.loc[1]} records ({target_pct.loc[1]:.2f}%). This level of imbalance means accuracy alone will not be a reliable evaluation metric.",
        "",
        "### Conclusion",
        "The target variable is heavily imbalanced and must be treated carefully during modeling.",
        "",
        "## Question 5",
        "### Aim",
        "Analyze numerical features.",
        "",
        "### Program",
        code_block(code_cells["q5_program"]),
        "",
        "### Output",
        q5_output,
        "",
        "### Analysis",
        "Age is broadly distributed and close to symmetric, while BMI and average glucose level are right-skewed. The boxplots and histograms show that glucose has the strongest upper tail and the clearest extreme values.",
        "",
        "### Conclusion",
        "The numerical variables are not normally distributed and will likely benefit from scaling and robust preprocessing.",
        "",
        "## Question 6",
        "### Aim",
        "Analyze categorical variables.",
        "",
        "### Program",
        code_block(code_cells["q6_program"]),
        "",
        "### Output",
        q6_output,
        "",
        "### Analysis",
        "Female records are more common than male records, Private work type is dominant, Urban and Rural residence are fairly balanced, and smoking_status includes a large Unknown group. The rare categories Other and Never_worked should be handled cautiously in later modeling.",
        "",
        "### Conclusion",
        "The categorical variables are usable, but a few rare classes may need consolidation or careful encoding.",
        "",
        "## Question 7",
        "### Aim",
        "Perform correlation analysis.",
        "",
        "### Program",
        code_block(code_cells["q7_program"]),
        "",
        "### Output",
        q7_output,
        "",
        "### Analysis",
        "The strongest positive relationships with stroke are age, heart disease, average glucose level, hypertension, and marital status. The strongest negative associations are with being unmarried, children, and unknown smoking status. The heatmap confirms that no single feature is perfectly correlated with the target, which is useful for a balanced modeling problem.",
        "",
        "### Conclusion",
        "Correlation signals are meaningful but moderate, so the final model will likely need multiple features rather than a single dominant predictor.",
        "",
        "## Question 8",
        "### Aim",
        "Perform outlier detection.",
        "",
        "### Program",
        code_block(code_cells["q8_program"]),
        "",
        "### Output",
        q8_output,
        "",
        "### Analysis",
        f"Outlier analysis shows {int(outliers.set_index('feature').loc['age', 'outlier_count'])} age outliers, {int(outliers.set_index('feature').loc['bmi', 'outlier_count'])} BMI outliers, and {int(outliers.set_index('feature').loc['avg_glucose_level', 'outlier_count'])} glucose outliers under the IQR rule. Average glucose level is the most outlier-prone variable.",
        "",
        "### Conclusion",
        "Outlier handling will be important, especially for average glucose level and BMI.",
        "",
        "## Final Analysis",
        *final_analysis,
        "",
        "## Final Conclusion",
        "The stroke prediction dataset is suitable for machine learning, but it is not ready for modeling without preprocessing. The main requirements are BMI imputation, categorical encoding, skewness handling, outlier awareness, and class-imbalance treatment. Once these steps are applied, the dataset should support reliable stroke-risk classification experiments.",
        "",
        "## Appendix: Generated Figures",
    ]

    for figure in get_figure_files():
        report.extend([
            f"### {figure.stem.replace('_', ' ').title()}",
            f"![{figure.stem}]({figure.relative_to(PROJECT_ROOT).as_posix()})",
            "",
        ])

    return "\n".join(report).rstrip() + "\n"


def set_docx_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
    styles["Title"].font.size = Pt(16)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.bold = True


def add_page_number_footer(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_markdown_image(document: Document, image_path: Path, width_inches: float = 6.2) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def add_code_block_docx(document: Document, code: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    run = paragraph.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def add_bullet_lines(document: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = document.add_paragraph(style="Normal")
        paragraph.style = document.styles["Normal"]
        paragraph.paragraph_format.left_indent = Inches(0.2)
        paragraph.style = document.styles["Normal"]
        run = paragraph.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def build_docx_report(markdown_report: str, df: pd.DataFrame, code_cells: dict[str, str], output_images: dict[str, Path], corr: pd.DataFrame) -> None:
    document = Document()
    set_docx_styles(document)
    add_page_number_footer(document)

    document.add_heading("Machine Learning Lab Report: Stroke Prediction Dataset", level=0)
    document.add_paragraph(
        "This lab report documents the exploratory analysis of the Kaggle stroke prediction dataset. The objective is to understand the dataset structure, quality, distributions, relationships, and preparation needs before building a classification model."
    )

    profile = dataframe_profile(df)
    missing = missing_summary(df)
    stats = statistical_summary(df)
    normality = normality_summary(df)
    outliers = outlier_summary(df)
    target_counts = target_distribution(df)
    target_pct = target_percentages(df)
    corr_table_md, _ = correlation_tables(df)

    sections = [
        ("Question 1", "Import, load and view the Stroke Prediction dataset.", code_cells["q1_program"], [output_images["q1_head"], output_images["q1_tail"], output_images["q1_shape_columns"], output_images["q1_dtypes"], output_images["q1_info"]], f"The dataset loads successfully from the workspace. It has {profile['shape'][0]} rows and {profile['shape'][1]} columns.", "The dataset is accessible and ready for analysis."),
        ("Question 2", "Display the summary statistics of the dataset.", code_cells["q2_program"], [output_images["q2_describe_numeric"], output_images["q2_describe_object"], output_images["q2_missing_values"], output_images["q2_duplicates"]], "The summary statistics show a wide age distribution, a right-skewed glucose feature, and missing BMI values.", "The descriptive statistics confirm that the dataset is realistic and suitable for EDA."),
        ("Question 3", "Analyze missing values and duplicate records.", code_cells["q3_program"], [(FIGURES_DIR / "missing_values_heatmap.png"), output_images["q2_missing_values"], output_images["q2_duplicates"]], f"Only BMI contains missing values, accounting for {int(profile['missing_values'].sum())} missing entries overall.", "The dataset requires missing-value handling for BMI, but duplicate removal is not necessary."),
    ]

    for title, aim, program, images, analysis, conclusion in sections:
        document.add_heading(title, level=1)
        document.add_paragraph("Aim", style="Heading 2")
        document.add_paragraph(aim)
        document.add_paragraph("Program", style="Heading 2")
        add_code_block_docx(document, program)
        document.add_paragraph("Output", style="Heading 2")
        for image_path in images:
            add_markdown_image(document, image_path)
        document.add_paragraph("Analysis", style="Heading 2")
        document.add_paragraph(analysis)
        document.add_paragraph("Conclusion", style="Heading 2")
        document.add_paragraph(conclusion)

    # Question 4
    q4_images = [FIGURES_DIR / "stroke_distribution.png", FIGURES_DIR / "stroke_pie_chart.png", FIGURES_DIR / "stroke_percentage_distribution.png"]
    document.add_heading("Question 4", level=1)
    document.add_paragraph("Aim", style="Heading 2")
    document.add_paragraph("Analyze the target variable (Stroke).")
    document.add_paragraph("Program", style="Heading 2")
    add_code_block_docx(document, code_cells["q4_program"])
    document.add_paragraph("Output", style="Heading 2")
    for image_path in q4_images:
        add_markdown_image(document, image_path)
    add_code_block_docx(document, markdown_table(["Stroke", "Count", "Percentage"], [[str(index), str(int(count)), f"{target_pct.loc[index]:.2f}%"] for index, count in target_counts.items()]))
    document.add_paragraph("Analysis", style="Heading 2")
    document.add_paragraph(f"Stroke cases are rare: class 0 has {target_counts.loc[0]} records ({target_pct.loc[0]:.2f}%) and class 1 has {target_counts.loc[1]} records ({target_pct.loc[1]:.2f}%).")
    document.add_paragraph("Conclusion", style="Heading 2")
    document.add_paragraph("The target variable is heavily imbalanced and must be treated carefully during modeling.")

    # Question 5
    q5_images = [FIGURES_DIR / "age_distribution.png", FIGURES_DIR / "bmi_distribution.png", FIGURES_DIR / "avg_glucose_level_distribution.png", FIGURES_DIR / "age_boxplot.png", FIGURES_DIR / "bmi_boxplot.png", FIGURES_DIR / "avg_glucose_level_boxplot.png"]
    document.add_heading("Question 5", level=1)
    document.add_paragraph("Aim", style="Heading 2")
    document.add_paragraph("Analyze numerical features.")
    document.add_paragraph("Program", style="Heading 2")
    add_code_block_docx(document, code_cells["q5_program"])
    document.add_paragraph("Output", style="Heading 2")
    for image_path in q5_images:
        add_markdown_image(document, image_path)
    document.add_paragraph("Analysis", style="Heading 2")
    document.add_paragraph("Age is broadly distributed and close to symmetric, while BMI and average glucose level are right-skewed. The boxplots show that glucose has the strongest upper tail and the clearest extreme values.")
    document.add_paragraph("Conclusion", style="Heading 2")
    document.add_paragraph("The numerical variables are not normally distributed and will likely benefit from scaling and robust preprocessing.")

    # Question 6
    q6_images = [FIGURES_DIR / "gender_countplot.png", FIGURES_DIR / "work_type_countplot.png", FIGURES_DIR / "ever_married_countplot.png", FIGURES_DIR / "Residence_type_countplot.png", FIGURES_DIR / "smoking_status_countplot.png", FIGURES_DIR / "hypertension_countplot.png", FIGURES_DIR / "heart_disease_countplot.png"]
    document.add_heading("Question 6", level=1)
    document.add_paragraph("Aim", style="Heading 2")
    document.add_paragraph("Analyze categorical variables.")
    document.add_paragraph("Program", style="Heading 2")
    add_code_block_docx(document, code_cells["q6_program"])
    document.add_paragraph("Output", style="Heading 2")
    for image_path in q6_images:
        add_markdown_image(document, image_path)
    document.add_paragraph("Analysis", style="Heading 2")
    document.add_paragraph("Female records are more common than male records, Private work type is dominant, Urban and Rural residence are fairly balanced, and smoking_status includes a large Unknown group.")
    document.add_paragraph("Conclusion", style="Heading 2")
    document.add_paragraph("The categorical variables are usable, but a few rare classes may need consolidation or careful encoding.")

    # Question 7
    document.add_heading("Question 7", level=1)
    document.add_paragraph("Aim", style="Heading 2")
    document.add_paragraph("Perform correlation analysis.")
    document.add_paragraph("Program", style="Heading 2")
    add_code_block_docx(document, code_cells["q7_program"])
    document.add_paragraph("Output", style="Heading 2")
    add_markdown_image(document, FIGURES_DIR / "correlation_heatmap.png")
    add_code_block_docx(document, corr.round(3).to_string())
    document.add_paragraph("Analysis", style="Heading 2")
    document.add_paragraph("The strongest positive relationships with stroke are age, heart disease, average glucose level, hypertension, and marital status. The strongest negative associations are with being unmarried, children, and unknown smoking status.")
    document.add_paragraph("Conclusion", style="Heading 2")
    document.add_paragraph("Correlation signals are meaningful but moderate, so the final model will likely need multiple features rather than a single dominant predictor.")

    # Question 8
    q8_images = [FIGURES_DIR / "age_outlier_boxplot.png", FIGURES_DIR / "bmi_outlier_boxplot.png", FIGURES_DIR / "avg_glucose_level_outlier_boxplot.png"]
    document.add_heading("Question 8", level=1)
    document.add_paragraph("Aim", style="Heading 2")
    document.add_paragraph("Perform outlier detection.")
    document.add_paragraph("Program", style="Heading 2")
    add_code_block_docx(document, code_cells["q8_program"])
    document.add_paragraph("Output", style="Heading 2")
    for image_path in q8_images:
        add_markdown_image(document, image_path)
    document.add_paragraph("Analysis", style="Heading 2")
    document.add_paragraph(f"Outlier analysis shows {int(outliers.set_index('feature').loc['age', 'outlier_count'])} age outliers, {int(outliers.set_index('feature').loc['bmi', 'outlier_count'])} BMI outliers, and {int(outliers.set_index('feature').loc['avg_glucose_level', 'outlier_count'])} glucose outliers under the IQR rule.")
    document.add_paragraph("Conclusion", style="Heading 2")
    document.add_paragraph("Outlier handling will be important, especially for average glucose level and BMI.")

    document.add_heading("Final Analysis", level=1)
    for bullet in [
        "Dataset size: 5,110 rows and 12 columns.",
        "Numerical features: age, avg_glucose_level, and bmi.",
        "Categorical features: gender, work_type, ever_married, Residence_type, and smoking_status.",
        f"Missing values: {int(profile['missing_values'].sum())}, concentrated in bmi.",
        f"Duplicate rows: {profile['duplicate_rows']}.",
        "Class imbalance: stroke is a strong minority class.",
        "Age is the strongest stroke-related feature.",
        "Heart disease and hypertension are positively associated with stroke.",
        "Average glucose level is right-skewed and outlier-prone.",
        "BMI is missing for 201 records and also shows moderate outlier behavior.",
        "Former smokers and self-employed patients have elevated stroke proportions.",
        "Residence type and gender show weak linear correlation with stroke.",
        "Overall data quality is good, with preprocessing still required before modeling.",
    ]:
        document.add_paragraph(bullet, style="List Bullet")

    document.add_heading("Final Conclusion", level=1)
    document.add_paragraph(
        "The stroke prediction dataset is suitable for machine learning, but it is not ready for modeling without preprocessing. The main requirements are BMI imputation, categorical encoding, skewness handling, outlier awareness, and class-imbalance treatment. Once these steps are applied, the dataset should support reliable stroke-risk classification experiments."
    )

    document.save(LAB_DOCX)


def register_report_fonts() -> None:
    font_dir = Path(r"C:\Windows\Fonts")
    pdfmetrics.registerFont(TTFont("Times New Roman", str(font_dir / "times.ttf")))
    pdfmetrics.registerFont(TTFont("Times New Roman Bold", str(font_dir / "timesbd.ttf")))
    pdfmetrics.registerFont(TTFont("Times New Roman Italic", str(font_dir / "timesi.ttf")))
    pdfmetrics.registerFont(TTFont("Times New Roman Bold Italic", str(font_dir / "timesbi.ttf")))


def build_pdf_report(markdown_report: str) -> None:
    register_report_fonts()

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="LabTitle",
            parent=styles["Title"],
            fontName="Times New Roman Bold",
            fontSize=16,
            leading=20,
            alignment=TA_CENTER,
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LabHeading1",
            parent=styles["Heading1"],
            fontName="Times New Roman Bold",
            fontSize=14,
            leading=18,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LabHeading2",
            parent=styles["Heading2"],
            fontName="Times New Roman Bold",
            fontSize=12,
            leading=16,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LabBody",
            parent=styles["BodyText"],
            fontName="Times New Roman",
            fontSize=12,
            leading=18,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LabCode",
            parent=styles["BodyText"],
            fontName="Courier",
            fontSize=8.5,
            leading=10,
            backColor=colors.whitesmoke,
            borderPadding=4,
            spaceAfter=8,
        )
    )

    def add_page_number(canvas, doc):
        canvas.setFont("Times New Roman", 10)
        canvas.drawRightString(A4[0] - 0.6 * inch, 0.4 * inch, f"Page {doc.page}")

    story = []
    lines = markdown_report.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# "):
            story.append(Paragraph(line[2:], styles["LabTitle"]))
        elif line.startswith("## "):
            story.append(Spacer(1, 0.08 * inch))
            story.append(Paragraph(line[3:], styles["LabHeading1"]))
        elif line.startswith("### "):
            story.append(Paragraph(line[4:], styles["LabHeading2"]))
        elif line.startswith("```python"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            story.append(Preformatted("\n".join(code_lines), styles["LabCode"]))
        elif line.startswith("!") and "(" in line and ")" in line:
            image_rel = line[line.find("(") + 1 : line.rfind(")")]
            image_path = PROJECT_ROOT / image_rel
            if image_path.exists():
                story.append(RLImage(str(image_path), width=6.6 * inch, height=4.2 * inch))
                story.append(Spacer(1, 0.08 * inch))
        elif line.startswith("| "):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].startswith("| "):
                table_lines.append(lines[i])
                i += 1
            rows = [[cell.strip() for cell in row.strip("|").split("|")] for row in table_lines]
            data = rows[0:1] + rows[2:] if len(rows) > 2 else rows
            table = Table(data, repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                        ("FONTNAME", (0, 0), (-1, -1), "Times New Roman"),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 0.08 * inch))
            continue
        elif line.strip().startswith("- "):
            story.append(Paragraph(line, styles["LabBody"]))
        elif line.strip():
            story.append(Paragraph(line, styles["LabBody"]))
        else:
            story.append(Spacer(1, 0.06 * inch))
        i += 1

    doc = SimpleDocTemplate(
        str(LAB_PDF),
        pagesize=A4,
        rightMargin=0.7 * inch,
        leftMargin=0.7 * inch,
        topMargin=0.7 * inch,
        bottomMargin=0.75 * inch,
    )
    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)


def main() -> None:
    ensure_directories()
    df = load_dataset(PROJECT_ROOT)
    create_missing_heatmap(df)
    code_cells = extract_notebook_cell_sources()
    output_images = create_text_output_images(df)
    _, corr = correlation_tables(df)
    markdown_report = build_markdown_report(df, code_cells, output_images, corr)
    LAB_MD.write_text(markdown_report, encoding="utf-8")
    build_docx_report(markdown_report, df, code_cells, output_images, corr)
    build_pdf_report(markdown_report)
    print(f"Markdown report written to: {LAB_MD}")
    print(f"DOCX report written to: {LAB_DOCX}")
    print(f"PDF report written to: {LAB_PDF}")
    print(f"Generated figures count: {len(get_figure_files())}")


if __name__ == "__main__":
    main()