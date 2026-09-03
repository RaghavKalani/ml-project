from __future__ import annotations

import json
import math
import textwrap
from dataclasses import dataclass
from io import StringIO
from pathlib import Path
from typing import Iterable

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from PIL import Image as PILImage
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as RLImage
from reportlab.platypus import PageBreak, Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle

from src.eda_utils import PROJECT_ROOT, encode_for_correlation, load_dataset, outlier_summary, normality_summary

NOTEBOOK_PATH = PROJECT_ROOT / "notebooks" / "01_EDA.ipynb"
REPORTS_DIR = PROJECT_ROOT / "reports"
FIGURES_DIR = PROJECT_ROOT / "figures"
MD_PATH = REPORTS_DIR / "ML_Lab_Report.md"
DOCX_PATH = REPORTS_DIR / "ML_Lab_Report.docx"
PDF_PATH = REPORTS_DIR / "ML_Lab_Report.pdf"

NUMERIC_FEATURES = ["age", "bmi", "avg_glucose_level"]
CATEGORICAL_FEATURES = ["gender", "work_type", "ever_married", "Residence_type", "smoking_status", "hypertension", "heart_disease"]
QUESTION1_OUTPUT_FILES = [
    "q1_head_output.png",
    "q1_tail_output.png",
    "q1_shape_columns_output.png",
    "q1_dtypes_output.png",
    "q1_info_output.png",
]
QUESTION2_OUTPUT_FILES = [
    "q2_describe_numeric_output.png",
    "q2_describe_object_output.png",
    "q2_missing_values_output.png",
    "q2_duplicate_count_output.png",
]
QUESTION4_FIGURES = ["stroke_distribution.png", "stroke_pie_chart.png", "stroke_percentage_distribution.png"]
QUESTION5_FIGURES = [
    "age_distribution.png",
    "bmi_distribution.png",
    "avg_glucose_level_distribution.png",
    "age_boxplot.png",
    "bmi_boxplot.png",
    "avg_glucose_level_boxplot.png",
    "age_violinplot.png",
    "bmi_violinplot.png",
    "avg_glucose_level_violinplot.png",
]
QUESTION6_FIGURES = [
    "gender_countplot.png",
    "work_type_countplot.png",
    "ever_married_countplot.png",
    "Residence_type_countplot.png",
    "smoking_status_countplot.png",
    "hypertension_countplot.png",
    "heart_disease_countplot.png",
]
QUESTION7_FIGURES = [
    "stroke_by_gender.png",
    "stroke_by_age_boxplot.png",
    "stroke_by_bmi_boxplot.png",
    "stroke_by_avg_glucose_level_boxplot.png",
    "stroke_by_smoking_status.png",
    "stroke_by_work_type.png",
    "stroke_by_residence_type.png",
    "stroke_by_hypertension.png",
    "stroke_by_heart_disease.png",
]
QUESTION8_FIGURES = ["correlation_matrix.png", "correlation_heatmap.png"]
QUESTION9_FIGURES = [
    "age_outlier_boxplot.png",
    "bmi_outlier_boxplot.png",
    "avg_glucose_level_outlier_boxplot.png",
    "qq_age.png",
    "qq_bmi.png",
    "qq_avg_glucose_level.png",
]


@dataclass
class FigureItem:
    path: Path
    caption: str


def ensure_directories() -> None:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    FIGURES_DIR.mkdir(parents=True, exist_ok=True)


def load_notebook_code_cells() -> list[str]:
    notebook = json.loads(NOTEBOOK_PATH.read_text(encoding="utf-8"))
    return ["".join(cell.get("source", [])) for cell in notebook["cells"] if cell.get("cell_type") == "code"]


def code_for_sections() -> dict[str, str]:
    code_cells = load_notebook_code_cells()
    return {
        "q1": "\n\n".join(code_cells[0:2]),
        "q2": code_cells[2],
        "q3": code_cells[3],
        "q4": code_cells[4],
        "q5": code_cells[5],
        "q6": code_cells[6],
        "q7": code_cells[7],
        "q8": code_cells[8],
        "q9": "\n\n".join(code_cells[9:11]),
    }


def ensure_eda_figures() -> None:
    expected = [
        *QUESTION1_OUTPUT_FILES,
        *QUESTION2_OUTPUT_FILES,
        "missing_values_heatmap.png",
        *QUESTION4_FIGURES,
        *QUESTION5_FIGURES,
        *QUESTION6_FIGURES,
        *QUESTION7_FIGURES,
        *QUESTION8_FIGURES,
        *QUESTION9_FIGURES,
    ]
    missing = [FIGURES_DIR / name for name in expected if not (FIGURES_DIR / name).exists()]
    if missing:
        from src.generate_eda import run as run_eda

        run_eda()


def render_text_image(text: str, output_name: str, title: str) -> Path:
    output_path = FIGURES_DIR / output_name
    if output_path.exists():
        return output_path

    lines = text.splitlines() or [""]
    height = max(3.0, 0.35 * len(lines) + 1.6)
    fig, ax = plt.subplots(figsize=(12, height))
    ax.axis("off")
    ax.set_title(title, fontsize=14, fontweight="bold", pad=12)
    ax.text(
        0.01,
        0.98,
        textwrap.fill(text, width=115),
        ha="left",
        va="top",
        family="monospace",
        fontsize=10,
        transform=ax.transAxes,
        bbox={"boxstyle": "round,pad=0.5", "facecolor": "white", "edgecolor": "#999999"},
    )
    fig.savefig(output_path, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return output_path


def dataframe_image(df: pd.DataFrame, output_name: str, title: str) -> Path:
    return render_text_image(df.to_string(), output_name, title)


def series_image(series: pd.Series, output_name: str, title: str) -> Path:
    return render_text_image(series.to_string(), output_name, title)


def unique_category_images(df: pd.DataFrame) -> list[Path]:
    images: list[Path] = []
    for column in ["gender", "work_type", "ever_married", "Residence_type", "smoking_status"]:
        counts = df[column].value_counts(dropna=False).to_frame(name="count")
        images.append(dataframe_image(counts, f"q3_unique_{column}.png", f"Unique values - {column}"))
    return images


def load_data() -> pd.DataFrame:
    return load_dataset(PROJECT_ROOT)


def make_correlation_matrix_image(df: pd.DataFrame) -> Path:
    output_path = FIGURES_DIR / "correlation_matrix.png"
    if output_path.exists():
        return output_path

    encoded = encode_for_correlation(df)
    corr = encoded.corr(numeric_only=True).round(2)
    fig, ax = plt.subplots(figsize=(max(12, len(corr.columns) * 0.6), max(10, len(corr.columns) * 0.6)))
    ax.axis("off")
    table = ax.table(cellText=corr.values, rowLabels=corr.index, colLabels=corr.columns, loc="center")
    table.auto_set_font_size(False)
    table.set_fontsize(6)
    table.scale(1.0, 1.2)
    fig.suptitle("Correlation Matrix", fontsize=14, fontweight="bold")
    fig.savefig(output_path, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return output_path


def q1_outputs(df: pd.DataFrame) -> list[FigureItem]:
    info_buf = StringIO()
    df.info(buf=info_buf)
    return [
        FigureItem(dataframe_image(df.head(), "q1_head_output.png", "Question 1 - First Five Rows"), "Figure 1. First five rows of the dataset."),
        FigureItem(dataframe_image(df.tail(), "q1_tail_output.png", "Question 1 - Last Five Rows"), "Figure 2. Last five rows of the dataset."),
        FigureItem(render_text_image(f"Shape: {df.shape}\n\nColumns: {list(df.columns)}", "q1_shape_columns_output.png", "Question 1 - Shape and Columns"), "Figure 3. Dataset shape and column names."),
        FigureItem(render_text_image(df.dtypes.to_string(), "q1_dtypes_output.png", "Question 1 - Data Types"), "Figure 4. Data types for each column."),
        FigureItem(render_text_image(info_buf.getvalue(), "q1_info_output.png", "Question 1 - info()"), "Figure 5. `info()` output for the dataset."),
    ]


def q2_outputs(df: pd.DataFrame) -> list[FigureItem]:
    missing = df.isna().sum().to_frame(name="missing_count")
    missing["missing_percentage"] = (missing["missing_count"] / len(df) * 100).round(2)
    return [
        FigureItem(dataframe_image(df.describe().round(3), "q2_describe_numeric_output.png", "Question 2 - describe()"), "Figure 6. Summary statistics for numerical columns."),
        FigureItem(dataframe_image(df.describe(include="object"), "q2_describe_object_output.png", "Question 2 - describe(include='object')"), "Figure 7. Summary statistics for categorical columns."),
        FigureItem(dataframe_image(missing, "q2_missing_values_output.png", "Question 2 - Missing Values"), "Figure 8. Missing value summary table."),
        FigureItem(render_text_image(f"Duplicate rows: {int(df.duplicated().sum())}", "q2_duplicate_count_output.png", "Question 2 - Duplicate Rows"), "Figure 9. Duplicate row count."),
    ]


def q3_outputs(df: pd.DataFrame) -> list[FigureItem]:
    missing = df.isna().sum().to_frame(name="missing_count")
    missing["missing_percentage"] = (missing["missing_count"] / len(df) * 100).round(2)
    missing_heatmap = FIGURES_DIR / "missing_values_heatmap.png"
    if not missing_heatmap.exists():
        from src.generate_eda import create_missing_heatmap

        create_missing_heatmap(df)
    items = [FigureItem(dataframe_image(missing, "q3_missing_table.png", "Question 3 - Missing Value Table"), "Figure 10. Missing value table."), FigureItem(missing_heatmap, "Figure 11. Missing value heatmap.")]
    for image in unique_category_images(df):
        figure_number = 12 + len(items) - 2
        items.append(FigureItem(image, f"Figure {figure_number}. Unique categorical values for {image.stem.replace('q3_unique_', '').replace('_', ' ')}."))
    return items


def q4_outputs() -> list[FigureItem]:
    target_dist = FIGURES_DIR / "stroke_distribution.png"
    pie = FIGURES_DIR / "stroke_pie_chart.png"
    perc = FIGURES_DIR / "stroke_percentage_distribution.png"
    return [
        FigureItem(target_dist, "Figure 17. Stroke count plot."),
        FigureItem(pie, "Figure 18. Stroke pie chart."),
        FigureItem(perc, "Table 1. Stroke class percentage distribution."),
    ]


def q5_outputs() -> list[FigureItem]:
    captions = [
        "Age histogram.",
        "BMI histogram.",
        "Glucose histogram.",
        "Age boxplot.",
        "BMI boxplot.",
        "Glucose boxplot.",
        "Age violin plot.",
        "BMI violin plot.",
        "Glucose violin plot.",
    ]
    names = QUESTION5_FIGURES
    return [FigureItem(FIGURES_DIR / name, f"Figure {20 + index}. {captions[index]}") for index, name in enumerate(names)]


def q6_outputs() -> list[FigureItem]:
    captions = [
        "Gender count plot.",
        "Work type count plot.",
        "Ever married count plot.",
        "Residence type count plot.",
        "Smoking status count plot.",
        "Hypertension count plot.",
        "Heart disease count plot.",
    ]
    return [FigureItem(FIGURES_DIR / name, f"Figure {29 + index}. {captions[index]}") for index, name in enumerate(QUESTION6_FIGURES)]


def q7_outputs() -> list[FigureItem]:
    captions = [
        "Stroke by gender.",
        "Stroke by age.",
        "Stroke by BMI.",
        "Stroke by glucose level.",
        "Stroke by smoking status.",
        "Stroke by work type.",
        "Stroke by residence type.",
        "Stroke by hypertension.",
        "Stroke by heart disease.",
    ]
    return [FigureItem(FIGURES_DIR / name, f"Figure {36 + index}. {captions[index]}") for index, name in enumerate(QUESTION7_FIGURES)]


def q8_outputs(df: pd.DataFrame) -> list[FigureItem]:
    correlation_matrix = make_correlation_matrix_image(df)
    heatmap = FIGURES_DIR / "correlation_heatmap.png"
    if not heatmap.exists():
        from src.generate_eda import create_correlation_heatmap

        create_correlation_heatmap(df)
    return [
        FigureItem(correlation_matrix, "Figure 45. Correlation matrix."),
        FigureItem(heatmap, "Figure 46. Correlation heatmap."),
    ]


def q9_outputs() -> list[FigureItem]:
    captions = [
        "Age outlier boxplot.",
        "BMI outlier boxplot.",
        "Glucose outlier boxplot.",
        "Age QQ plot.",
        "BMI QQ plot.",
        "Glucose QQ plot.",
    ]
    return [FigureItem(FIGURES_DIR / name, f"Figure {47 + index}. {captions[index]}") for index, name in enumerate(QUESTION9_FIGURES)]


def markdown_table(headers: list[str], rows: list[list[str]]) -> str:
    head = "| " + " | ".join(headers) + " |"
    sep = "| " + " | ".join(["---"] * len(headers)) + " |"
    body = ["| " + " | ".join(row) + " |" for row in rows]
    return "\n".join([head, sep, *body])


def summary_tables(df: pd.DataFrame) -> tuple[str, str, str, str]:
    profile = {
        "shape": f"{df.shape[0]} rows x {df.shape[1]} columns",
        "missing": int(df.isna().sum().sum()),
        "duplicates": int(df.duplicated().sum()),
        "numerical": len(df.select_dtypes(include="number").columns) - 2,  # exclude id and stroke? keep dataset features only
        "categorical": len(df.select_dtypes(include="object").columns) + 2,
    }
    # dataset-level tables used in the final analysis section
    target_counts = df["stroke"].value_counts().sort_index()
    target_pct = (target_counts / len(df) * 100).round(2)
    target_table = markdown_table(["Stroke", "Count", "Percentage"], [[str(i), str(int(target_counts.loc[i])), f"{target_pct.loc[i]:.2f}%"] for i in target_counts.index])
    missing = df.isna().sum().to_frame(name="Missing Count")
    missing["Missing %"] = (missing["Missing Count"] / len(df) * 100).round(2)
    missing_table = missing.to_markdown()
    stat = outlier_summary(df)
    normal = normality_summary(df)
    return target_table, missing_table, stat.to_markdown(index=False), normal.to_markdown(index=False)


def build_report_content(df: pd.DataFrame, code: dict[str, str]) -> tuple[str, dict[str, str], list[FigureItem]]:
    target_counts = df["stroke"].value_counts().sort_index()
    target_pct = (target_counts / len(df) * 100).round(2)
    missing = df.isna().sum().to_frame(name="Missing Count")
    missing["Missing %"] = (missing["Missing Count"] / len(df) * 100).round(2)
    duplicates = int(df.duplicated().sum())
    corr = encode_for_correlation(df).corr(numeric_only=True).round(3)
    corr_strength = corr["stroke"].sort_values(ascending=False)
    top_pos = corr_strength.iloc[1:6]
    top_neg = corr_strength.tail(5).sort_values()
    outliers = outlier_summary(df)
    normality = normality_summary(df)
    num = df[NUMERIC_FEATURES].copy()
    stat_rows = []
    for column in NUMERIC_FEATURES:
        s = num[column].dropna()
        stat_rows.append(
            [
                column,
                f"{s.mean():.4f}",
                f"{s.median():.4f}",
                f"{s.mode().iloc[0]:.4f}",
                f"{s.var():.4f}",
                f"{s.std():.4f}",
                f"{s.skew():.4f}",
                f"{s.kurt():.4f}",
            ]
        )
    stats_table = markdown_table(["Feature", "Mean", "Median", "Mode", "Variance", "Std Dev", "Skewness", "Kurtosis"], stat_rows)

    q1_analysis = [
        f"- The dataset loads correctly with {df.shape[0]} rows and {df.shape[1]} columns.",
        f"- The file includes the expected stroke-related variables and mixed data types.",
        f"- The shape, column names, and memory usage confirm that the data is small enough for manual EDA.",
        f"- No loading issues appear in the notebook outputs.",
    ]
    q2_analysis = [
        f"- The numerical summary shows that age ranges from {df['age'].min():.2f} to {df['age'].max():.2f}.",
        f"- Average glucose level is strongly right-skewed, while BMI is moderately right-skewed.",
        f"- `describe(include='object')` confirms that Female, Private, Yes, Urban, and never smoked are the dominant categories.",
        f"- The data types remain consistent with a structured tabular dataset.",
    ]
    q3_analysis = [
        f"- BMI is the only column with missing values, and it contributes {int(df['bmi'].isna().sum())} missing records.",
        f"- The heatmap shows a very narrow missingness pattern rather than scattered nulls.",
        f"- Duplicate rows are absent, which simplifies downstream preprocessing.",
        f"- The categorical variables have a few rare levels such as Other and Never_worked.",
    ]
    q4_analysis = [
        f"- Stroke class 0 contains {int(target_counts.loc[0])} records ({target_pct.loc[0]:.2f}%).",
        f"- Stroke class 1 contains {int(target_counts.loc[1])} records ({target_pct.loc[1]:.2f}%).",
        "- The class distribution is highly imbalanced and accuracy alone would be misleading.",
        "- The pie chart and count plot both show that stroke is a minority class.",
    ]
    q5_analysis = [
        "- Age is broadly spread and visually close to a symmetric distribution.",
        "- BMI is right-skewed with a long upper tail.",
        "- Average glucose level has the strongest skewness and the most extreme values.",
        "- The violin plots confirm that the central mass of BMI and glucose is concentrated below the upper tails.",
    ]
    q6_analysis = [
        "- Female records are more common than male records, with a single Other record.",
        "- Private work type is the largest category, followed by Self-employed and children.",
        "- Residence type is nearly balanced between Urban and Rural.",
        "- Smoking status includes a large Unknown group, which should be handled carefully later.",
    ]
    q7_analysis = [
        "- Stroke is visibly more common among older patients.",
        "- Patients with hypertension or heart disease show a higher stroke proportion.",
        "- Former smokers and self-employed patients show elevated stroke proportions.",
        "- Residence type has a much weaker relationship with stroke than the clinical factors.",
    ]
    q8_analysis = [
        "- Age has the strongest positive correlation with stroke among the encoded variables.",
        "- Heart disease, average glucose level, hypertension, and marital status also show positive association.",
        "- Being unmarried and being in the children work type are negatively correlated with stroke.",
        "- The heatmap shows moderate, not extreme, relationships overall.",
    ]
    q9_analysis = [
        f"- Outlier counts under the IQR rule are {int(outliers.set_index('feature').loc['age', 'outlier_count'])} for age, {int(outliers.set_index('feature').loc['bmi', 'outlier_count'])} for BMI, and {int(outliers.set_index('feature').loc['avg_glucose_level', 'outlier_count'])} for glucose.",
        "- Average glucose level has the clearest outlier spread.",
        "- The QQ plots confirm that all three numerical variables deviate from normality.",
        f"- Normality testing marks age as {bool(normality.loc[normality['feature'] == 'age', 'normal_at_0.05'].iloc[0])}, BMI as {bool(normality.loc[normality['feature'] == 'bmi', 'normal_at_0.05'].iloc[0])}, and glucose as {bool(normality.loc[normality['feature'] == 'avg_glucose_level', 'normal_at_0.05'].iloc[0])} at the 0.05 level.",
    ]

    target_table = markdown_table(["Stroke", "Count", "Percentage"], [[str(i), str(int(target_counts.loc[i])), f"{target_pct.loc[i]:.2f}%"] for i in target_counts.index])
    missing_table = markdown_table(["Column", "Missing Count", "Missing %"], [[index, str(int(row['Missing Count'])), f"{row['Missing %']:.2f}%"] for index, row in missing.iterrows()])
    unique_tables = []
    for column in ["gender", "work_type", "ever_married", "Residence_type", "smoking_status"]:
        counts = df[column].value_counts(dropna=False)
        unique_tables.append((column, markdown_table([column, "Count"], [[str(index), str(int(value))] for index, value in counts.items()])))
    corr_table = markdown_table(["Feature", "Correlation with Stroke"], [[index, f"{value:.3f}"] for index, value in corr_strength.items()])
    outlier_table = markdown_table(["Feature", "Lower Bound", "Upper Bound", "Outlier Count", "Outlier %"], [[row['feature'], f"{row['lower_bound']:.4f}", f"{row['upper_bound']:.4f}", str(int(row['outlier_count'])), f"{row['outlier_percentage']:.2f}%"] for _, row in outliers.iterrows()])
    normal_table = markdown_table(["Feature", "Normaltest Statistic", "P-value", "Normal at 0.05"], [[row['feature'], f"{row['normaltest_statistic']:.4f}", f"{row['p_value']:.6f}", str(bool(row['normal_at_0.05']))] for _, row in normality.iterrows()])

    figures: list[FigureItem] = []
    figures.extend(q1_outputs(df))
    figures.extend(q2_outputs(df))
    # question 3 uses the table images and the heatmap plus the unique-category images
    figures.extend(q3_outputs(df))
    figures.extend(q4_outputs())
    figures.extend(q5_outputs())
    figures.extend(q6_outputs())
    figures.extend(q7_outputs())
    figures.extend(q8_outputs(df))
    figures.extend(q9_outputs())

    markdown_sections = [
        "# Machine Learning Laboratory",
        "",
        "Experiment:",
        "Acute Ischemic Stroke Risk Classification Engine – Dataset Exploration & Exploratory Data Analysis",
        "",
        "Dataset:",
        "Healthcare Stroke Prediction Dataset",
        "",
        "--------------------------------------------------",
        "",
    ]

    def section(question: int, aim: str, program: str, output: str, analysis_lines: list[str], conclusion: str) -> list[str]:
        return [
            f"Question {question}",
            "",
            "Aim",
            aim,
            "",
            "Program",
            "```python",
            program.rstrip(),
            "```",
            "",
            "Output",
            output,
            "",
            "Analysis",
            *analysis_lines,
            "",
            "Conclusion",
            conclusion,
            "",
            "--------------------------------------------------",
            "",
        ]

    markdown_sections.extend(section(1, "Import, load and view the dataset.", code["q1"], "\n\n".join([f"![]({item.path.relative_to(PROJECT_ROOT).as_posix()})\n\n{item.caption}" for item in q1_outputs(df)]), q1_analysis, "The dataset loads successfully and is ready for EDA."))
    markdown_sections.extend(section(2, "Display dataset overview and summary statistics.", code["q2"], "\n\n".join([f"![]({item.path.relative_to(PROJECT_ROOT).as_posix()})\n\n{item.caption}" for item in q2_outputs(df)]), q2_analysis, "The summary statistics show a realistic clinical dataset with skewness and missing BMI values."))
    q3_items = q3_outputs(df)
    q3_output_text = [f"![]({image.path.relative_to(PROJECT_ROOT).as_posix()})\n\n{image.caption}" for image in q3_items]
    markdown_sections.extend(section(3, "Analyze data quality.", code["q3"], "\n\n".join(q3_output_text), q3_analysis, "The data quality is good overall, with BMI missingness as the main issue."))
    markdown_sections.extend(section(4, "Analyze the target variable.", code["q4"], "\n\n".join([f"![]({item.path.relative_to(PROJECT_ROOT).as_posix()})\n\n{item.caption}" for item in q4_outputs()]) + f"\n\n{target_table}", q4_analysis, "Stroke is a severely imbalanced target variable."))
    markdown_sections.extend(section(5, "Analyze numerical features.", code["q5"], "\n\n".join([f"![]({item.path.relative_to(PROJECT_ROOT).as_posix()})\n\n{item.caption}" for item in q5_outputs()]), q5_analysis, "The numerical variables are skewed and require preprocessing."))
    markdown_sections.extend(section(6, "Analyze categorical variables.", code["q6"], "\n\n".join([f"![]({item.path.relative_to(PROJECT_ROOT).as_posix()})\n\n{item.caption}" for item in q6_outputs()]), q6_analysis, "The categorical variables are clean but contain a few rare levels."))
    markdown_sections.extend(section(7, "Perform bivariate analysis.", code["q7"], "\n\n".join([f"![]({item.path.relative_to(PROJECT_ROOT).as_posix()})\n\n{item.caption}" for item in q7_outputs()]), q7_analysis, "Older age, hypertension, heart disease, and glucose are the clearest bivariate signals."))
    markdown_sections.extend(section(8, "Perform correlation analysis.", code["q8"], "\n\n".join([f"![]({item.path.relative_to(PROJECT_ROOT).as_posix()})\n\n{item.caption}" for item in q8_outputs(df)]) + f"\n\n{corr_table}", q8_analysis, "Age is the strongest positive correlation with stroke."))
    markdown_sections.extend(section(9, "Perform outlier detection and distribution analysis.", code["q9"], "\n\n".join([f"![]({item.path.relative_to(PROJECT_ROOT).as_posix()})\n\n{item.caption}" for item in q9_outputs()]) + f"\n\n{outlier_table}\n\n{normal_table}", q9_analysis, "The numerical features are non-normal and outlier-prone."))

    markdown_sections.extend([
        "Analysis",
        "- Dataset dimensions: 5,110 rows and 12 columns.",
        "- Missing values: 201 missing BMI values.",
        f"- Duplicate rows: {duplicates}.",
        "- Numerical features: age, bmi, avg_glucose_level.",
        "- Categorical features: gender, work_type, ever_married, Residence_type, smoking_status.",
        f"- Target imbalance: stroke class 0 dominates at {target_pct.loc[0]:.2f}% while class 1 is only {target_pct.loc[1]:.2f}%.",
        "- Age has the strongest relationship with stroke.",
        "- Glucose and heart disease also show clear positive association with stroke.",
        "- BMI is missing for a subset of rows and is moderately skewed.",
        "- Hypertension and heart disease are important clinical signals.",
        "- Outliers are most visible in glucose, followed by BMI.",
        "- Correlations are moderate rather than extreme, which is useful for classification.",
        "- The numerical variables are not normally distributed.",
        "- Preprocessing should include imputation, encoding, and imbalance handling.",
        "",
        "Conclusion",
        "The stroke prediction dataset is suitable for machine learning, but preprocessing is required before modeling. The EDA shows a strongly imbalanced target, missing BMI values, right-skewed numerical variables, and clinically meaningful relationships with age, glucose level, hypertension, and heart disease. These patterns suggest that the dataset can support a classification model, but only after imputation, encoding, scaling, and imbalance-aware training are applied. The current dataset quality is good enough to proceed to the next phase, provided that the preprocessing pipeline is carefully designed.",
        "",
    ])

    used_figures = {item.path.resolve() for item in figures}
    remaining_figures = sorted(path for path in FIGURES_DIR.glob("*.png") if path.resolve() not in used_figures)
    if remaining_figures:
        markdown_sections.extend([
            "Additional Figures",
            "",
        ])
        for index, figure_path in enumerate(remaining_figures, start=1):
            markdown_sections.extend([
                f"Figure A{index}. {figure_path.stem.replace('_', ' ').title()}.",
                f"![]({figure_path.relative_to(PROJECT_ROOT).as_posix()})",
                "",
            ])

    return "\n".join(markdown_sections), code, figures


def add_page_numbers(canvas, doc):
    canvas.setFont("Times-Roman", 10)
    canvas.drawRightString(A4[0] - 0.6 * inch, 0.4 * inch, f"Page {doc.page}")


def register_pdf_fonts() -> None:
    font_dir = Path(r"C:\Windows\Fonts")
    pdfmetrics.registerFont(TTFont("Times-Roman", str(font_dir / "times.ttf")))
    pdfmetrics.registerFont(TTFont("Times-Bold", str(font_dir / "timesbd.ttf")))
    pdfmetrics.registerFont(TTFont("Times-Italic", str(font_dir / "timesi.ttf")))
    pdfmetrics.registerFont(TTFont("Times-BoldItalic", str(font_dir / "timesbi.ttf")))


def image_size(path: Path, max_width: float = 6.4 * inch) -> tuple[float, float]:
    with PILImage.open(path) as img:
        width, height = img.size
    scale = min(1.0, max_width / float(width))
    return width * scale, height * scale


def build_docx(markdown: str) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.space_after = Pt(6)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        document.styles[style_name].font.name = "Times New Roman"

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)

    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
        elif line.startswith("```python"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            p = document.add_paragraph()
            r = p.add_run("\n".join(code_lines))
            r.font.name = "Courier New"
            r.font.size = Pt(9)
        elif line.startswith("![]("):
            path = PROJECT_ROOT / line[line.find("(") + 1 : line.find(")")]
            if path.exists():
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(path), width=Inches(6.4))
        elif line.startswith("| "):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].startswith("| "):
                table_lines.append(lines[i])
                i += 1
            rows = [[c.strip() for c in row.strip("|").split("|")] for row in table_lines]
            if len(rows) > 2:
                rows = [rows[0], *rows[2:]]
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx, value in enumerate(row):
                    table.cell(r_idx, c_idx).text = value
        elif line.startswith("-"):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.strip():
            document.add_paragraph(line.strip())
        i += 1

    document.save(DOCX_PATH)


def build_pdf(markdown: str) -> None:
    register_pdf_fonts()
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="LabTitle", fontName="Times-Bold", fontSize=16, leading=20, alignment=TA_CENTER, spaceAfter=10))
    styles.add(ParagraphStyle(name="LabHeading1", fontName="Times-Bold", fontSize=14, leading=18, spaceAfter=8))
    styles.add(ParagraphStyle(name="LabHeading2", fontName="Times-Bold", fontSize=12, leading=16, spaceAfter=6))
    styles.add(ParagraphStyle(name="LabBody", fontName="Times-Roman", fontSize=12, leading=18, spaceAfter=5))
    styles.add(ParagraphStyle(name="LabCode", fontName="Courier", fontSize=9, leading=10, spaceAfter=8))

    story = []
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# "):
            story.append(Paragraph(line[2:].strip(), styles["LabTitle"]))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:].strip(), styles["LabHeading1"]))
        elif line.startswith("### "):
            story.append(Paragraph(line[4:].strip(), styles["LabHeading2"]))
        elif line.startswith("```python"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            story.append(Preformatted("\n".join(code_lines), styles["LabCode"]))
        elif line.startswith("![]("):
            image_rel = line[line.find("(") + 1 : line.find(")")]
            image_path = PROJECT_ROOT / image_rel
            if image_path.exists():
                width, height = image_size(image_path)
                story.append(RLImage(str(image_path), width=min(width, 6.3 * inch), height=min(height, 8.0 * inch)))
                story.append(Spacer(1, 0.08 * inch))
        elif line.startswith("| "):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].startswith("| "):
                table_lines.append(lines[i])
                i += 1
            rows = [[c.strip() for c in row.strip("|").split("|")] for row in table_lines]
            if len(rows) > 2:
                rows = [rows[0], *rows[2:]]
            table = Table(rows, repeatRows=1)
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("FONTNAME", (0, 0), (-1, -1), "Times-Roman"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
            ]))
            story.append(table)
            story.append(Spacer(1, 0.08 * inch))
        elif line.startswith("-"):
            story.append(Paragraph(line[2:].strip(), styles["LabBody"]))
        elif line.strip():
            story.append(Paragraph(line.strip(), styles["LabBody"]))
        else:
            story.append(Spacer(1, 0.06 * inch))
        i += 1

    doc = SimpleDocTemplate(str(PDF_PATH), pagesize=A4, rightMargin=0.7 * inch, leftMargin=0.7 * inch, topMargin=0.7 * inch, bottomMargin=0.75 * inch)
    doc.build(story, onFirstPage=add_page_numbers, onLaterPages=add_page_numbers)


def main() -> None:
    ensure_directories()
    ensure_eda_figures()
    df = load_data()
    code = code_for_sections()
    markdown, _, figures = build_report_content(df, code)
    MD_PATH.write_text(markdown, encoding="utf-8")
    build_docx(markdown)
    build_pdf(markdown)
    print(f"Wrote {MD_PATH}")
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {PDF_PATH}")
    print(f"Figures referenced: {len(figures)}")


if __name__ == "__main__":
    main()